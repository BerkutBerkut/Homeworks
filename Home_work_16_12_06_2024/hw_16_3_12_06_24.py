"""
Домашнее задание №11: Работа с комплексными файлами - excel, json, 
word. Git и Jira
а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран.
в) Создайте абзац с текстом и запишите это в новый word-файл, изменит шрифт и размер 
шрифта.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# Создание нового документа
doc = Document()

# Добавление текста
paragraph = doc.add_paragraph()
run = paragraph.add_run("Hello Python, ")

run = paragraph.add_run("Pyton Hello!")
run.bold = True

# Сохранение документа
doc.save("doc_file.docx")

# Открытие существующего документа
doc = Document("doc_file.docx")

# Поиск и чтение жирного текста
bold_text = ""
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:  # Проверка, является ли текст жирным
            bold_text += run.text

# Вывод жирного текста на экран
print("Жирный текст:", bold_text)

# Создание нового документа
new_doc = Document()

# Добавление абзаца с текстом
paragraph = new_doc.add_paragraph("This is Home Work of PYTHON.")

# Настройка шрифта и размера
run = paragraph.runs[0]
run.font.name = "Algerian"  # Установка шрифта
run.font.size = Pt(24)  # Установка размера шрифта

# Сохранение нового документа
new_doc.save("new_file.docx")
